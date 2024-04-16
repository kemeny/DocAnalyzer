import os
from dotenv import load_dotenv
import streamlit as st
import re
from docx import Document
from openai import AzureOpenAI
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

load_dotenv()

# Initialize OpenAI client
def initialize_clients():
    client = AzureOpenAI(
        azure_endpoint = "https://ai-testinghubresourceeus2794140685701.openai.azure.com/", 
        api_key=os.getenv("AZURE_OPENAI_KEY"),  
        api_version="2024-02-15-preview"
    )
    return client

def parse_user_suggestions(input_string):
    try:
        # Assuming user inputs JSON-like suggestions: [{'find':'term1', 'replace':'term2'}, ...]
        suggestions = eval(input_string)
        assert isinstance(suggestions, list)
        for item in suggestions:
            assert 'find' in item and 'replace' in item
        return suggestions
    except:
        st.error("Invalid format for suggestions. Please input as [{'find':'term1', 'replace':'term2'}, ...]")
        return []

def filter_content(text, client, user_suggestions):
    # Check and apply user suggestions to filter the content
    if user_suggestions:
        for index, suggestion in enumerate(user_suggestions):
            if not isinstance(suggestion, dict) or 'find' not in suggestion or 'replace' not in suggestion:
                raise ValueError(f"Suggestion at index {index} is invalid. Expected a dictionary with 'find' and 'replace', got: {suggestion}")
            text = text.replace(suggestion['find'], suggestion['replace'])
    
    # Find all unique client names in the text
    client_names = set(re.findall(r'\bClient Name \d+\b', text))
    name_map = {}

    for client_name in client_names:
        if client_name not in name_map:
            prompt = f"Generate a generic name for the client name '{client_name}':"
            try:
                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "system", "content": prompt}],
                    temperature=0.7,
                    max_tokens=60
                )
                if response.choices:
                    generic_name = response.choices[0].message.content
                    name_map[client_name] = generic_name
                else:
                    raise ValueError(f"No valid response for client name {client_name}")
            except Exception as e:
                print(f"Error fetching generic name for {client_name}: {e}")
                name_map[client_name] = client_name  # Use original name if API call fails

        text = text.replace(client_name, name_map[client_name])

    # Filter out any inappropriate content
    text = re.sub(r'\b(?:sex|sexual|violence|hate)\b', '***', text, flags=re.IGNORECASE)

    return text



def read_docx(file):
    doc = Document(file)
    full_text = ' '.join(para.text for para in doc.paragraphs if para.text.strip())  # Join all text into a single string
    return full_text


def analyze_sequence(uploaded_files, client):
    analysis_report = []
    proposed_emails = []

    # Process each file and generate reports
    emails = [read_docx(file) for file in uploaded_files]  # Each email is now a single string
    for i, email_content in enumerate(emails):
        email_filtered = filter_content(email_content, client.azure_client, client.user_suggestions)
        try:
            sentiment_response = client.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "system", "content": email_filtered}],
                temperature=0.7,
                max_tokens=60
            )
            sentiment_score = sentiment_response.choices[0].message.content if sentiment_response.choices else None
            if sentiment_score is None:
                raise ValueError("Failed to get sentiment score")
            grade = 'A' if 'positive' in sentiment_score else 'B' if 'neutral' in sentiment_score else 'C'
            suggestion = generate_style_suggestions(client, email_filtered)
            analysis_report.append(f"Email {i+1}: {suggestion}. Sentiment: {sentiment_score}. Grade: {grade}")
            proposed_emails.append(suggestion)
        except Exception as e:
            analysis_report.append(f"Email {i+1}: Error processing this email - {e}")

    return analysis_report, proposed_emails



# Generate suggestions for style alignment
def generate_style_suggestions(client, text):
    message_text = [{"role":"system","content": text}]
    response = client.chat.completions.create(
        model="gpt-4", # model = "deployment_name"
        messages = message_text,
        temperature=0.7,
        max_tokens=800,
        top_p=0.95,
        frequency_penalty=0,
        presence_penalty=0,
        stop=None
    )
    return response.choices[0].message.content.strip()

# Generate DOCX file with proposed emails
def generate_docx(emails):
    doc = Document()
    para = doc.add_paragraph()
    for email in emails:
        para.add_run(email + '\n\n')
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para.space_after = Pt(14)
    doc.save("proposed_emails.docx")
    return "proposed_emails.docx"

# Streamlit app layout
st.title('Document Analyzer')

uploaded_files = st.file_uploader(
    "Upload DOC files containing the emails you want to analyze", 
    accept_multiple_files=True, 
    help="Each DOC file should contain one email. The emails will be analyzed for style and a revised version will be proposed."
)

azure_client = initialize_clients()

if uploaded_files:
    st.success("Files uploaded successfully!")
    emails = [read_docx(file) for file in uploaded_files]
    st.write("Emails:", emails)
    user_suggestions_input = st.text_input("Enter your suggestions for filtering the content:")
    user_suggestions = parse_user_suggestions(user_suggestions_input)
    if st.button("Apply suggestions"):
        emails = [filter_content(email, azure_client, user_suggestions) for email in emails]
        st.write("Filtered Emails:", emails)
        if st.button("Approve filtered content"):
            report, proposed_emails = analyze_sequence(emails, azure_client)
            st.write("Analysis Report:", report)
            user_suggestions_input = st.text_input(
                "Enter your suggestions for filtering the content:",
                help="Input as [{'find':'term1', 'replace':'term2'}, ...]"
            )
            if st.button("Apply suggestions"):
                proposed_emails = [generate_style_suggestions(azure_client, email, user_suggestions) for email in proposed_emails]
                st.write("Proposed New Emails:", proposed_emails)
                if st.button("Approve proposed emails"):
                    docx_file = generate_docx(proposed_emails)
                    st.download_button("Download the revised emails", docx_file)